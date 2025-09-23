# file_processor.py - Multi-format document processing

import os
import tempfile
from typing import Dict, List
from PIL import Image
import pytesseract

# File type support
try:
    import docx
    DOCX_AVAILABLE = True
    print("python-docx loaded successfully")
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Word document support disabled.")

# OCR functionality
try:
    import easyocr
    EASYOCR_AVAILABLE = True
    print("EasyOCR loaded successfully")
except ImportError:
    EASYOCR_AVAILABLE = False
    print("Warning: EasyOCR not available. Image OCR functionality disabled.")

def extract_text_from_docx(file_path: str) -> Dict:
    """
    Extract text from a Word document (.docx)
    Returns a dict with pages, quality_assessment, total_characters, readable_pages
    """
    if not DOCX_AVAILABLE:
        return {
            "pages": [],
            "quality_assessment": "error",
            "total_characters": 0,
            "readable_pages": 0,
            "total_pages": 0,
            "quality_issues": ["Word document support not available"]
        }
    
    try:
        doc = docx.Document(file_path)
        
        # Extract text from all paragraphs
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Split into pages (approximate - Word doesn't have exact page breaks)
        # We'll treat each paragraph as a "page" for consistency
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        pages = []
        total_characters = len(full_text)
        readable_pages = len(paragraphs)
        
        for i, paragraph_text in enumerate(paragraphs, start=1):
            char_count = len(paragraph_text)
            pages.append({
                "page_number": i,
                "text": paragraph_text,
                "quality": "good",  # Word documents are usually good quality
                "char_count": char_count,
                "extraction_method": "docx"
            })
        
        return {
            "pages": pages,
            "quality_assessment": "good",
            "total_characters": total_characters,
            "readable_pages": readable_pages,
            "total_pages": len(paragraphs),
            "quality_issues": []
        }
        
    except Exception as e:
        return {
            "pages": [],
            "quality_assessment": "error",
            "total_characters": 0,
            "readable_pages": 0,
            "total_pages": 0,
            "quality_issues": [f"Error reading Word document: {str(e)}"]
        }

def extract_text_from_image(file_path: str) -> Dict:
    """
    Extract text from an image file (.jpg, .png, .tiff, etc.)
    Returns a dict with pages, quality_assessment, total_characters, readable_pages
    """
    try:
        # Open image
        image = Image.open(file_path)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Extract text using OCR
        if EASYOCR_AVAILABLE:
            # Use EasyOCR for better accuracy
            reader = easyocr.Reader(['en'])
            results = reader.readtext(image)
            
            # Combine all detected text
            extracted_text = ""
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # Only include high-confidence text
                    extracted_text += text + " "
            
            extracted_text = extracted_text.strip()
        else:
            # Fallback to pytesseract
            extracted_text = pytesseract.image_to_string(image)
        
        if not extracted_text or len(extracted_text.strip()) < 10:
            return {
                "pages": [],
                "quality_assessment": "unreadable",
                "total_characters": 0,
                "readable_pages": 0,
                "total_pages": 0,
                "quality_issues": ["No readable text found in image"]
            }
        
        # Create a single "page" for the image
        pages = [{
            "page_number": 1,
            "text": extracted_text,
            "quality": "fair",  # OCR text is usually fair quality
            "char_count": len(extracted_text),
            "extraction_method": "ocr"
        }]
        
        return {
            "pages": pages,
            "quality_assessment": "fair",
            "total_characters": len(extracted_text),
            "readable_pages": 1,
            "total_pages": 1,
            "quality_issues": []
        }
        
    except Exception as e:
        return {
            "pages": [],
            "quality_assessment": "error",
            "total_characters": 0,
            "readable_pages": 0,
            "total_pages": 0,
            "quality_issues": [f"Error processing image: {str(e)}"]
        }

def extract_text_from_file(file_path: str, file_type: str) -> Dict:
    """
    Extract text from various file types
    Returns a dict with pages, quality_assessment, total_characters, readable_pages
    """
    file_type = file_type.lower()
    
    if file_type == "pdf":
        # Import here to avoid circular imports
        from pdf_parser import extract_text_from_pdf
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type in ["jpg", "jpeg", "png", "tiff", "bmp", "gif"]:
        return extract_text_from_image(file_path)
    else:
        return {
            "pages": [],
            "quality_assessment": "error",
            "total_characters": 0,
            "readable_pages": 0,
            "total_pages": 0,
            "quality_issues": [f"Unsupported file type: {file_type}"]
        }

def get_supported_file_types() -> List[str]:
    """
    Get list of supported file types
    """
    supported_types = ["pdf"]
    
    if DOCX_AVAILABLE:
        supported_types.extend(["docx"])
    
    if EASYOCR_AVAILABLE or pytesseract:
        supported_types.extend(["jpg", "jpeg", "png", "tiff", "bmp", "gif"])
    
    return supported_types

def validate_file_type(filename: str) -> tuple:
    """
    Validate if file type is supported
    Returns (is_supported, file_type, error_message)
    """
    if not filename:
        return False, None, "No filename provided"
    
    # Get file extension
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ""
    
    if not file_ext:
        return False, None, "No file extension found"
    
    supported_types = get_supported_file_types()
    
    if file_ext not in supported_types:
        return False, file_ext, f"Unsupported file type: .{file_ext}. Supported types: {', '.join(supported_types)}"
    
    return True, file_ext, None
